"""Utilitaires pour l'export en DOCX avec styles configurables."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re
import os
from datetime import datetime
from utils.docx_style_loader import get_style_config, load_style_config

# Chemin par défaut vers le fichier de configuration
DEFAULT_STYLE_CONFIG_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    'config',
    'docx_styles.yml'
)


def create_element(name):
    """Crée un élément XML pour les propriétés avancées."""
    return OxmlElement(name)


def add_page_number(paragraph):
    """Ajoute un numéro de page au paragraphe."""
    run = paragraph.add_run()

    # Texte avant le numéro
    run.text = "Document rédigé par une IA - Veuillez vérifier le contenu - Page "

    # Numéro de page actuel
    fldChar1 = create_element('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = create_element('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Texte " / "
    run = paragraph.add_run(" / ")

    # Nombre total de pages
    fldChar3 = create_element('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')

    instrText2 = create_element('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar3)
    run._r.append(instrText2)
    run._r.append(fldChar4)


def _get_default_style_config():
    """Charge la configuration de style depuis /config/docx_styles.yml."""
    if not os.path.exists(DEFAULT_STYLE_CONFIG_PATH):
        raise FileNotFoundError(
            f"Le fichier de configuration des styles est introuvable : {DEFAULT_STYLE_CONFIG_PATH}"
        )
    return load_style_config(DEFAULT_STYLE_CONFIG_PATH)


def add_formatted_text(paragraph, text, style_config=None):
    """
    Ajoute du texte avec formatage inline (gras, italique, code).

    Args:
        paragraph: Paragraphe docx
        text: Texte à formater
        style_config: Configuration de styles (optionnel)
    """
    if style_config is None:
        style_config = _get_default_style_config()

    # Pattern pour détecter **gras**, *italique*, `code`
    # Amélioration : gérer ** avant * pour éviter les conflits
    pattern = r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        # Code inline (priorité haute)
        if part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            code_style = style_config.styles.get('code_inline', {}).get('font', {})

            if 'name' in code_style:
                run.font.name = code_style['name']
            if 'size' in code_style:
                run.font.size = Pt(code_style['size'])

            if 'color' in code_style:
                color_hex = code_style['color']
                if color_hex.startswith('#'):
                    color_hex = color_hex[1:]
                try:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except (ValueError, IndexError):
                    raise ValueError(f"Couleur invalide dans la configuration : {code_style['color']}")

        # Gras
        elif part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True

        # Italique
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True

        # Texte normal
        else:
            paragraph.add_run(part)


def add_markdown_table(doc, table_lines, style_config=None):
    """
    Ajoute un tableau markdown au document avec style sobre.

    Args:
        doc: Document docx
        table_lines: Lignes du tableau markdown
        style_config: Configuration de styles (optionnel)
    """
    if style_config is None:
        style_config = _get_default_style_config()

    if not table_lines:
        return

    # Parser les lignes
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')]
        # Retirer les cellules vides au début et à la fin
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Créer le tableau
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)

    # Utiliser le style de tableau depuis la configuration
    table.style = style_config.get_table_style_name()

    # Remplir le tableau
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]

                # Effacer le texte par défaut
                cell.text = ''

                # Traiter les sauts de ligne <br> et <br/>
                cell_parts = re.split(r'<br\s*/?>', cell_text)

                for part_idx, part in enumerate(cell_parts):
                    if part_idx > 0:
                        # Ajouter un nouveau paragraphe pour chaque <br>
                        para = cell.add_paragraph()
                    else:
                        # Utiliser le paragraphe existant
                        para = cell.paragraphs[0]

                    # Appliquer le formatage Markdown au texte
                    add_formatted_text(para, part.strip(), style_config)

                    # En-tête en gras (sauf si déjà en gras via **)
                    if i == 0:
                        for run in para.runs:
                            if not run.bold:
                                run.bold = True

    doc.add_paragraph()

#####################################################################################
#    Crée un document DOCX contenant une question et sa réponse.
#
#    Args:
#        question: Question posée
#        response: Réponse générée
#        metadata: Métadonnées optionnelles (date, modèle, etc.)
#        style_config_path: Chemin vers le fichier de configuration de styles
#
#    Returns:
#        BytesIO contenant le document DOCX
#####################################################################################
def create_response_docx(
    question: str,
    response: str,
    metadata: dict = None,
    style_config_path: str = None
) -> BytesIO:

    # Charger la configuration de styles
    if style_config_path:
        style_config = load_style_config(style_config_path)
    else:
        # Utiliser le fichier de configuration par défaut /config/docx_styles.yml
        style_config = _get_default_style_config()

    doc = Document()

    # Métadonnées du document depuis la config
    config_metadata = style_config.get_metadata()
    doc_title = config_metadata['title']
    doc_author = config_metadata['author']

    # En-tête du document
    header = doc.add_heading(doc_title, level=1)
    style_config.apply_heading_style(header, level=1)

    # Métadonnées
    if metadata:
        p = doc.add_paragraph()
        p.add_run('Date : ').bold = True
        p.add_run(metadata.get('date', datetime.now().strftime('%d/%m/%Y %H:%M')))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if metadata.get('model'):
            p = doc.add_paragraph()
            p.add_run('Modèle : ').bold = True
            p.add_run(metadata['model'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()

    # Convertir la réponse Markdown en DOCX
    response_lines = response.split('\n')
    i = 0

    while i < len(response_lines):
        line = response_lines[i]

        # Titres
        if line.startswith('#### '):
            h = doc.add_heading(line[5:].strip(), level=4)
            style_config.apply_heading_style(h, level=4)
            i += 1
            continue
        elif line.startswith('### '):
            h = doc.add_heading(line[4:].strip(), level=3)
            style_config.apply_heading_style(h, level=3)
            i += 1
            continue
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=3)
            style_config.apply_heading_style(h, level=3)
            i += 1
            continue
        elif line.startswith('# '):
            h = doc.add_heading(line[2:].strip(), level=2)
            style_config.apply_heading_style(h, level=2)
            i += 1
            continue

        # Listes à puces
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text, style_config)
            style_config.apply_list_bullet_style(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            i += 1
            continue

        # Listes numérotées
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text, style_config)
            style_config.apply_list_number_style(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            i += 1
            continue

        # Tableaux markdown
        elif '|' in line and i + 1 < len(response_lines) and '---' in response_lines[i + 1]:
            table_lines = [line]
            i += 2  # Sauter la ligne de séparation
            while i < len(response_lines) and '|' in response_lines[i]:
                table_lines.append(response_lines[i])
                i += 1

            # Créer le tableau
            add_markdown_table(doc, table_lines, style_config)
            continue

        # Blocs de code
        elif line.strip().startswith('```'):
            i += 1
            code_lines = []
            lang = line.strip()[3:].strip()  # Langue du code
            while i < len(response_lines) and not response_lines[i].strip().startswith('```'):
                code_lines.append(response_lines[i])
                i += 1
            i += 1  # Sauter le ``` de fermeture

            # Ajouter le code avec le style configuré
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                style_config.apply_code_style(p)
            continue

        # Séparateurs horizontaux
        elif line.strip() in ['---', '***', '___']:
            doc.add_paragraph('─' * 50)
            i += 1
            continue

        # Ligne vide
        elif not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Paragraphe normal
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line, style_config)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            i += 1

    # Pied de page avec numérotation
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Ajouter le numéro de page
    add_page_number(footer_para)

    # Style du pied de page
    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    # Sauvegarder
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
